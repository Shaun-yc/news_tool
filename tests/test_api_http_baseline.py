import hashlib
import io
import json
import os
import tempfile
import unittest
import zipfile
from datetime import datetime
from pathlib import Path
from unittest.mock import patch

import openpyxl
from docx import Document
from fastapi.testclient import TestClient

from api import app


class VllmResponse:
    def raise_for_status(self):
        return None

    def json(self):
        return {"choices": [{"message": {"content": "NONE"}}]}


class ApiHttpBaselineTests(unittest.TestCase):
    def test_process_rejects_corrupt_docx_upload(self):
        with TestClient(app, raise_server_exceptions=False) as client:
            response = client.post(
                "/process",
                files={
                    "file": (
                        "broken.docx",
                        b"not-a-word-package",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"], "請上傳有效的 .docx 週新聞檔案")

    def test_process_rejects_non_word_ooxml_package(self):
        package = io.BytesIO()
        with zipfile.ZipFile(package, "w") as archive:
            archive.writestr("readme.txt", "not a word document")

        with TestClient(app, raise_server_exceptions=False) as client:
            response = client.post(
                "/process",
                files={
                    "file": (
                        "not-word.docx",
                        package.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"], "請上傳有效的 .docx 週新聞檔案")

    def test_process_rejects_malformed_word_ooxml_package(self):
        document = Document()
        document.add_paragraph("A valid document package.")
        docx_file = io.BytesIO()
        document.save(docx_file)
        valid_package = docx_file.getvalue()
        malformed_members = (
            "[Content_Types].xml",
            "_rels/.rels",
            "word/document.xml",
        )

        with TestClient(app, raise_server_exceptions=False) as client:
            for member in malformed_members:
                package = io.BytesIO()
                with zipfile.ZipFile(io.BytesIO(valid_package), "r") as source:
                    with zipfile.ZipFile(package, "w") as target:
                        for info in source.infolist():
                            data = b"<root/>" if info.filename == member else source.read(info.filename)
                            target.writestr(info, data)

                with self.subTest(member=member):
                    response = client.post(
                        "/process",
                        files={
                            "file": (
                                "malformed.docx",
                                package.getvalue(),
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        },
                    )

                    self.assertEqual(response.status_code, 400)
                    self.assertEqual(response.json()["detail"], "請上傳有效的 .docx 週新聞檔案")

    def test_process_rejects_malformed_content_types_xml_reader_errors(self):
        document = Document()
        document.add_paragraph("A valid document package.")
        docx_file = io.BytesIO()
        document.save(docx_file)
        valid_package = docx_file.getvalue()
        malformed_payloads = (
            b"<root>",
            b'<?xml version="1.0" encoding="x-unknown"?><root/>',
        )

        with TestClient(app, raise_server_exceptions=False) as client:
            for payload in malformed_payloads:
                package = io.BytesIO()
                with zipfile.ZipFile(io.BytesIO(valid_package), "r") as source:
                    with zipfile.ZipFile(package, "w") as target:
                        for info in source.infolist():
                            data = (
                                payload
                                if info.filename == "[Content_Types].xml"
                                else source.read(info.filename)
                            )
                            target.writestr(info, data)

                with self.subTest(payload=payload):
                    response = client.post(
                        "/process",
                        files={
                            "file": (
                                "malformed-content-types.docx",
                                package.getvalue(),
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        },
                    )

                    self.assertEqual(response.status_code, 400)
                    self.assertEqual(response.json()["detail"], "請上傳有效的 .docx 週新聞檔案")

    def test_process_preserves_expired_non_audit_directory(self):
        document = Document()
        document.add_paragraph(",Manual backup preservation")
        document.add_paragraph("A single news item without a source URL.")
        docx_file = io.BytesIO()
        document.save(docx_file)
        docx_bytes = docx_file.getvalue()

        with tempfile.TemporaryDirectory() as temp_dir:
            manual_backup = Path(temp_dir) / "manual_backup"
            manual_backup.mkdir()
            backup_file = manual_backup / "keep.txt"
            backup_file.write_text("do not delete", encoding="utf-8")
            expired_time = datetime.now().timestamp() - 31 * 24 * 60 * 60
            os.utime(manual_backup, (expired_time, expired_time))

            environment = {
                "VLLM_BASE_URL": "http://summary.local",
                "VLLM_MODEL": "summary-model",
                "VLLM_TIMEOUT_SECONDS": "30",
                "VLLM_TEMPERATURE": "0",
                "VLLM_MAX_TOKENS": "256",
                "CLASSIFY_BASE_URL": "http://classify.local",
                "CLASSIFY_MODEL": "classify-model",
                "CLASSIFY_MAX_TOKENS": "64",
                "SUMMARY_ALIGN_MAX_TOKENS": "384",
                "SCRAPE_DELAY_SECONDS": "0",
                "CLASSIFY_DELAY_SECONDS": "0",
                "REQUEST_TIMEOUT_SECONDS": "7",
                "AUDIT_ARCHIVE_DIR": temp_dir,
                "AUDIT_RETENTION_DAYS": "30",
            }
            with patch.dict(os.environ, environment, clear=False):
                with patch("services.classifier.requests.post", return_value=VllmResponse()):
                    with TestClient(app) as client:
                        response = client.post(
                            "/process",
                            files={
                                "file": (
                                    "weekly.2026.08.13.docx",
                                    docx_bytes,
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                )
                            },
                        )

            self.assertEqual(response.status_code, 200)
            self.assertTrue(manual_backup.exists())
            self.assertEqual(backup_file.read_text(encoding="utf-8"), "do not delete")

    def test_process_returns_report_headers_workbook_and_audit_archive(self):
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "項次"
        table.cell(0, 1).text = "標題"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "API端到端新聞"
        document.add_paragraph("API端到端新聞")
        document.add_paragraph("API 真實 HTTP 測試摘要。")
        docx_file = io.BytesIO()
        document.save(docx_file)
        docx_bytes = docx_file.getvalue()
        filename = "週新聞_2026.08.13.docx"

        expected_headers = [
            "doc_id",
            "headline",
            "responsible_unit",
            "compiled_by",
            "project",
            "functional_cat",
            "subcategory",
            "pubdate",
            "source",
            "source_url",
            "content",
            "content_tran",
            "attach_cnt",
            "attach_fname",
        ]
        expected_values = [
            [
                "20260813_01",
                "API端到端新聞",
                "碳費推動組",
                "永智顧問",
                "碳市場國際合作與企業能力建構計畫",
                "排放管理;國際事務;調適韌性;減量交易",
                "待人工確認",
                None,
                "LINK",
                None,
                "API 真實 HTTP 測試摘要。",
                "來源資料沒有網址，請手動確認原文。",
                0,
                None,
            ]
        ]

        with tempfile.TemporaryDirectory() as temp_dir:
            environment = {
                "VLLM_BASE_URL": "http://summary.local",
                "VLLM_MODEL": "summary-model",
                "VLLM_TIMEOUT_SECONDS": "30",
                "VLLM_TEMPERATURE": "0",
                "VLLM_MAX_TOKENS": "256",
                "CLASSIFY_BASE_URL": "http://classify.local",
                "CLASSIFY_MODEL": "classify-model",
                "CLASSIFY_MAX_TOKENS": "64",
                "SUMMARY_ALIGN_MAX_TOKENS": "384",
                "SCRAPE_DELAY_SECONDS": "0",
                "CLASSIFY_DELAY_SECONDS": "0",
                "REQUEST_TIMEOUT_SECONDS": "7",
                "AUDIT_ARCHIVE_DIR": temp_dir,
                "AUDIT_RETENTION_DAYS": "30",
            }
            with patch.dict(os.environ, environment, clear=False):
                with patch("services.classifier.requests.post", return_value=VllmResponse()):
                    with TestClient(app) as client:
                        response = client.post(
                            "/process",
                            files={
                                "file": (
                                    filename,
                                    docx_bytes,
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                )
                            },
                        )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(
                response.headers["content-type"],
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            self.assertEqual(
                response.headers["content-disposition"],
                "attachment; filename=weekly-news.xlsx; "
                "filename*=UTF-8''%E6%B0%B8%E6%99%BA%E9%80%B1%E6%96%B0%E8%81%9Ecsv_20260813.xlsx",
            )
            self.assertEqual(response.headers["x-news-total-count"], "1")
            self.assertEqual(response.headers["x-news-scrape-failed-count"], "1")
            self.assertEqual(response.headers["x-news-classification-fallback-count"], "1")
            self.assertEqual(response.headers["x-news-summary-aligned-count"], "0")
            workbook = openpyxl.load_workbook(io.BytesIO(response.content))
            worksheet = workbook.active
            self.assertEqual(worksheet.title, "週新聞")
            self.assertEqual(worksheet.max_column, 14)
            self.assertEqual([cell.value for cell in worksheet[1]], expected_headers)
            self.assertEqual(
                [
                    [cell.value for cell in row]
                    for row in worksheet.iter_rows(min_row=2, max_row=2, max_col=14)
                ],
                expected_values,
            )

            audit_root = Path(temp_dir)
            audit_dirs = [path for path in audit_root.iterdir() if path.is_dir()]
            self.assertEqual(len(audit_dirs), 1)
            audit_dir = audit_dirs[0]
            self.assertRegex(audit_dir.name, r"^\d{8}T\d{6}\.\d{6}Z_[0-9a-f]{12}$")
            self.assertEqual((audit_dir / "input.docx").read_bytes(), docx_bytes)
            self.assertEqual((audit_dir / "output.xlsx").read_bytes(), response.content)

            metadata = json.loads((audit_dir / "metadata.json").read_text(encoding="utf-8"))
            self.assertEqual(
                set(metadata),
                {
                    "audit_id",
                    "created_at",
                    "input_filename",
                    "output_filename",
                    "input_sha256",
                    "total_count",
                    "scrape_failed_count",
                    "classification_fallback_count",
                    "summary_aligned_count",
                },
            )
            self.assertEqual(metadata["audit_id"], audit_dir.name)
            self.assertEqual(metadata["input_filename"], filename)
            self.assertEqual(metadata["output_filename"], "永智週新聞csv_20260813.xlsx")
            self.assertEqual(metadata["input_sha256"], hashlib.sha256(docx_bytes).hexdigest())
            self.assertEqual(metadata["total_count"], 1)
            self.assertEqual(metadata["scrape_failed_count"], 1)
            self.assertEqual(metadata["classification_fallback_count"], 1)
            self.assertEqual(metadata["summary_aligned_count"], 0)
            created_at = datetime.fromisoformat(metadata["created_at"])
            self.assertIsNotNone(created_at.tzinfo)
            self.assertTrue(audit_dir.name.startswith(created_at.strftime("%Y%m%dT%H%M%S.")))
            self.assertEqual(audit_dir.name.rsplit("_", 1)[1], metadata["input_sha256"][:12])


if __name__ == "__main__":
    unittest.main()
