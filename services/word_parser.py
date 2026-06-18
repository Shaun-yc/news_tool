import re

from docx import Document


CATEGORY_TITLES = [
    "國際間碳市場發展",
    "國內外碳市場最新動態",
    "國內碳市場發展",
]


def parse_word_news(docx_file_object):
    """Parse the summary table and article paragraphs from an uploaded Word file."""
    doc = Document(docx_file_object)
    titles_from_table = []

    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            if any("項次" in text or "標題" in text for text in cells_text):
                continue
            if len(cells_text) < 2:
                continue

            title = cells_text[1].strip().strip(",\"' \n\r")
            if title and not any(category in title for category in CATEGORY_TITLES):
                if title not in titles_from_table:
                    titles_from_table.append(title)

    if not titles_from_table:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith(",") and len(text) > 5:
                title = text.strip(", ")
                if not any(category in title for category in CATEGORY_TITLES):
                    titles_from_table.append(title)

    paragraphs_text = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    news_list = []

    for target_title in titles_from_table:
        content_segments = []
        source_url = ""

        for index, text in enumerate(paragraphs_text):
            if text != target_title and not (len(target_title) > 10 and target_title[:10] in text):
                continue

            for following_text in paragraphs_text[index + 1 :]:
                if "出處" in following_text or "Source" in following_text or "http" in following_text:
                    url_match = re.search(r"https?://[^\s]+", following_text)
                    if url_match:
                        source_url = url_match.group(0)
                    break
                if following_text in titles_from_table:
                    break
                content_segments.append(following_text)
            break

        full_content = "\n".join(content_segments)
        news_list.append(
            {
                "zh_title": target_title,
                "content": full_content or target_title,
                "source_url": source_url,
            }
        )

    return news_list
