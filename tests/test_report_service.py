import io
import json
import unittest
from unittest.mock import patch

import openpyxl
from docx import Document

from services.config import Settings
from services.report_service import build_weekly_news_report


SUCCESS_URL = "https://93.184.216.34/news/success"
CLASSIFY_URL = "http://classify.local"
SUMMARY_URL = "http://summary.local"
SUCCESS_BODY = (
    "This source article provides a verified carbon market update with policy details "
    "and implementation context for the weekly report."
)
ALIGNMENT_RESPONSE = (
    "政府公布碳費制度最新執行規則，說明企業申報、查核與減量措施，並透過監督機制提升制度透明度與市場可預測性。"
    "政府同時要求企業依據法定程序提交排放資料，主管機關將持續檢視執行情形並提供必要的技術指引。"
    "此項制度也鼓勵企業投入減量技術，並以透明的監測結果支持後續碳市場合作。"
)


def _success_html():
    structured_data = {
        "@type": "NewsArticle",
        "headline": "Successful source headline",
        "datePublished": "2026-06-01T08:00:00Z",
        "articleBody": SUCCESS_BODY,
    }
    return (
        "<html><head><script type='application/ld+json'>"
        f"{json.dumps(structured_data, ensure_ascii=False)}"
        "</script></head><body></body></html>"
    )


class HtmlResponse:
    is_redirect = False
    headers = {}
    apparent_encoding = "utf-8"
    encoding = None

    def __init__(self, text):
        self.text = text

    def raise_for_status(self):
        return None


class DeterministicSession:
    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_value, traceback):
        return None

    def get(self, url, headers, timeout, allow_redirects):
        if url == SUCCESS_URL:
            return HtmlResponse(_success_html())
        raise AssertionError(f"unexpected source URL: {url}")


class VllmResponse:
    def __init__(self, content):
        self._content = content

    def raise_for_status(self):
        return None

    def json(self):
        return {"choices": [{"message": {"content": self._content}}]}


class ReportServiceTests(unittest.TestCase):
    def test_build_weekly_news_report_keeps_end_to_end_report_and_processing_summary(self):
        document = Document()
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "項次"
        table.cell(0, 1).text = "標題"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "成功新聞甲篇"
        table.cell(2, 0).text = "2"
        table.cell(2, 1).text = "失敗新聞乙篇"
        document.add_paragraph("成功新聞甲篇")
        document.add_paragraph("第一篇原始摘要。")
        document.add_paragraph(f"出處：{SUCCESS_URL}")
        document.add_paragraph("失敗新聞乙篇")
        document.add_paragraph("第二篇原始摘要。")
        file_object = io.BytesIO()
        document.save(file_object)
        file_object.seek(0)

        settings = Settings(
            vllm_base_url=SUMMARY_URL,
            vllm_model="summary-model",
            vllm_temperature=0,
            vllm_max_tokens=256,
            scrape_delay_seconds=0,
            classify_delay_seconds=0,
            request_timeout_seconds=7,
            vllm_timeout_seconds=30,
            classify_base_url=CLASSIFY_URL,
            classify_model="classify-model",
            classify_max_tokens=64,
            summary_align_max_tokens=384,
        )

        post_calls = []

        def fake_post(url, **kwargs):
            post_calls.append(url)
            if url == f"{CLASSIFY_URL}/v1/chat/completions":
                if post_calls.count(url) == 1:
                    return VllmResponse("碳定價;氣候法制")
                return VllmResponse("NONE")
            if url == f"{SUMMARY_URL}/v1/chat/completions":
                return VllmResponse(ALIGNMENT_RESPONSE)
            raise AssertionError(f"unexpected vLLM URL: {url}")

        with patch("services.classifier.requests.post", side_effect=fake_post):
            report, output_filename, summary = build_weekly_news_report(
                file_object,
                "weekly.2026.06.05.docx",
                settings,
                session_factory=DeterministicSession,
                sleep=lambda seconds: None,
            )

        self.assertEqual(output_filename, "永智週新聞csv_20260605.xlsx")
        self.assertEqual(summary.total_count, 2)
        self.assertEqual(summary.scrape_failed_count, 1)
        self.assertEqual(summary.classification_fallback_count, 1)
        self.assertEqual(summary.summary_aligned_count, 1)

        worksheet = openpyxl.load_workbook(report).active
        self.assertEqual(
            [
                [cell.value for cell in row]
                for row in worksheet.iter_rows(min_row=2, max_row=3, max_col=14)
            ],
            [
                [
                    "20260605_01",
                    "成功新聞甲篇\nSuccessful source headline",
                    "碳費推動組",
                    "永智顧問",
                    "碳市場國際合作與企業能力建構計畫",
                    "排放管理;國際事務;調適韌性;減量交易",
                    "碳定價;氣候法制",
                    "2026/06/01",
                    "93",
                    SUCCESS_URL,
                    ALIGNMENT_RESPONSE,
                    SUCCESS_BODY,
                    0,
                    None,
                ],
                [
                    "20260605_02",
                    "失敗新聞乙篇",
                    "碳費推動組",
                    "永智顧問",
                    "碳市場國際合作與企業能力建構計畫",
                    "排放管理;國際事務;調適韌性;減量交易",
                    "待人工確認",
                    None,
                    "LINK",
                    None,
                    "第二篇原始摘要。",
                    "來源資料沒有網址，請手動確認原文。",
                    0,
                    None,
                ],
            ],
        )


if __name__ == "__main__":
    unittest.main()
