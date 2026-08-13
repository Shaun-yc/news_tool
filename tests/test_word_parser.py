import io
import unittest

from docx import Document

from services.word_parser import parse_word_news


class WordParserTests(unittest.TestCase):
    @staticmethod
    def _as_docx_file(document):
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def test_parse_word_news_reads_table_title_content_and_source(self):
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "項次"
        table.cell(0, 1).text = "標題"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "歐盟宣布碳市場新措施"
        document.add_paragraph("歐盟宣布碳市場新措施")
        document.add_paragraph("這是一段新聞摘要。")
        document.add_paragraph("出處：https://www.example.com/news/1")
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        news_list = parse_word_news(buffer)

        self.assertEqual(len(news_list), 1)
        self.assertEqual(news_list[0]["zh_title"], "歐盟宣布碳市場新措施")
        self.assertEqual(news_list[0]["content"], "這是一段新聞摘要。")
        self.assertEqual(news_list[0]["source_url"], "https://www.example.com/news/1")

    def test_parse_word_news_keeps_each_article_summary_and_source_for_two_news(self):
        document = Document()
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "項次"
        table.cell(0, 1).text = "標題"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "台灣碳費政策觀察報告甲"
        table.cell(2, 0).text = "2"
        table.cell(2, 1).text = "歐盟碳市場制度觀察報告乙"
        document.add_paragraph("台灣碳費政策觀察報告甲")
        document.add_paragraph("台灣碳費第一篇摘要。")
        document.add_paragraph("出處：https://www.example.com/news/taiwan")
        document.add_paragraph("歐盟碳市場制度觀察報告乙")
        document.add_paragraph("歐盟碳市場第二篇摘要。")
        document.add_paragraph("Source: https://www.example.com/news/eu")

        news_list = parse_word_news(self._as_docx_file(document))

        self.assertEqual(
            news_list,
            [
                {
                    "zh_title": "台灣碳費政策觀察報告甲",
                    "content": "台灣碳費第一篇摘要。",
                    "source_url": "https://www.example.com/news/taiwan",
                },
                {
                    "zh_title": "歐盟碳市場制度觀察報告乙",
                    "content": "歐盟碳市場第二篇摘要。",
                    "source_url": "https://www.example.com/news/eu",
                },
            ],
        )

    def test_parse_word_news_uses_current_prefix_match_for_titles_sharing_first_ten_chars(self):
        document = Document()
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "項次"
        table.cell(0, 1).text = "標題"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "能源政策共同前綴測試甲新聞"
        table.cell(2, 0).text = "2"
        table.cell(2, 1).text = "能源政策共同前綴測試乙新聞"
        document.add_paragraph("能源政策共同前綴測試甲新聞")
        document.add_paragraph("共享前綴測試的第一篇摘要。")
        document.add_paragraph("出處：https://www.example.com/news/prefix-first")
        document.add_paragraph("能源政策共同前綴測試乙新聞")
        document.add_paragraph("共享前綴測試的第二篇摘要。")
        document.add_paragraph("出處：https://www.example.com/news/prefix-second")

        news_list = parse_word_news(self._as_docx_file(document))

        self.assertEqual(
            news_list,
            [
                {
                    "zh_title": "能源政策共同前綴測試甲新聞",
                    "content": "共享前綴測試的第一篇摘要。",
                    "source_url": "https://www.example.com/news/prefix-first",
                },
                {
                    "zh_title": "能源政策共同前綴測試乙新聞",
                    "content": "共享前綴測試的第一篇摘要。",
                    "source_url": "https://www.example.com/news/prefix-first",
                },
            ],
        )

    def test_parse_word_news_keeps_summary_and_empty_source_when_no_url_exists(self):
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "項次"
        table.cell(0, 1).text = "標題"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "沒有來源網址的新聞案例"
        document.add_paragraph("沒有來源網址的新聞案例")
        document.add_paragraph("這篇新聞沒有提供來源網址。")

        news_list = parse_word_news(self._as_docx_file(document))

        self.assertEqual(
            news_list,
            [
                {
                    "zh_title": "沒有來源網址的新聞案例",
                    "content": "這篇新聞沒有提供來源網址。",
                    "source_url": "",
                }
            ],
        )

    def test_parse_word_news_uses_comma_prefixed_paragraph_as_title_without_table(self):
        document = Document()
        document.add_paragraph(",備援新聞來源測試案例甲")
        document.add_paragraph("沒有表格時的備援摘要。")
        document.add_paragraph("出處：https://www.example.com/news/fallback")

        news_list = parse_word_news(self._as_docx_file(document))

        self.assertEqual(
            news_list,
            [
                {
                    "zh_title": "備援新聞來源測試案例甲",
                    "content": "沒有表格時的備援摘要。",
                    "source_url": "https://www.example.com/news/fallback",
                }
            ],
        )


if __name__ == "__main__":
    unittest.main()

